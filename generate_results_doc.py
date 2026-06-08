"""Generate AC_PINN_Results.docx - a complete printout of all 7 notebooks plus
metrics/ablation/noise-study/training tables. Local-only artifact, not committed."""

import os
import sys
import json
import base64
import csv
import io

import numpy as np
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_PATH = 'AC_PINN_Results.docx'

NOTEBOOKS = [
    ('notebooks/00_setup_and_verify.ipynb', 'Notebook 00: Setup and Verify'),
    ('notebooks/01_burgers.ipynb',          'Notebook 01: Burgers Equation'),
    ('notebooks/02_heat.ipynb',             'Notebook 02: Heat Equation'),
    ('notebooks/03_wave.ipynb',             'Notebook 03: Wave Equation'),
    ('notebooks/04_allen_cahn.ipynb',       'Notebook 04: Allen-Cahn Equation'),
    ('notebooks/05_ablation.ipynb',         'Notebook 05: Ablation Study'),
    ('notebooks/06_final_comparison.ipynb', 'Notebook 06: Final Comparison'),
]

PDES = ['burgers', 'heat', 'wave', 'allen_cahn']
PDE_LABELS = {'burgers': 'Burgers', 'heat': 'Heat', 'wave': 'Wave', 'allen_cahn': 'Allen-Cahn'}

CODE_FONT = 'Courier New'
CODE_SIZE = Pt(9)
TEXT_FONT = 'Arial'
TEXT_SIZE = Pt(11)


# ----------------------------------------------------------------------------
# Helpers for adding content to the document
# ----------------------------------------------------------------------------

def add_code_block(doc, text):
    if not text.strip():
        return
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = CODE_FONT
    run.font.size = CODE_SIZE
    p.paragraph_format.space_after = Pt(2)


def add_output_text(doc, text):
    if not text.strip():
        return
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = CODE_FONT
    run.font.size = CODE_SIZE
    run.font.color.rgb = None
    p.paragraph_format.space_after = Pt(2)


def add_normal_text(doc, text):
    if not text.strip():
        return
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = TEXT_FONT
    run.font.size = TEXT_SIZE


def add_image_from_bytes(doc, img_bytes, caption=None):
    try:
        stream = io.BytesIO(img_bytes)
        doc.add_picture(stream, width=Inches(6))
        if caption:
            cap = doc.add_paragraph()
            run = cap.add_run(caption)
            run.font.name = TEXT_FONT
            run.font.size = Pt(9)
            run.italic = True
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        add_normal_text(doc, f'[Could not insert image: {e}]')


def add_image_from_disk(doc, path, caption=None):
    try:
        doc.add_picture(path, width=Inches(6))
        if caption:
            cap = doc.add_paragraph()
            run = cap.add_run(caption)
            run.font.name = TEXT_FONT
            run.font.size = Pt(9)
            run.italic = True
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        add_normal_text(doc, f'[Could not load image {path}: {e}]')


def make_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        run = hdr_cells[i].paragraphs[0].add_run(str(h))
        run.bold = True
        run.font.name = TEXT_FONT
        run.font.size = TEXT_SIZE
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.name = TEXT_FONT
            run.font.size = Pt(10)
    return table


# ----------------------------------------------------------------------------
# Notebook processing
# ----------------------------------------------------------------------------

def process_markdown_cell(doc, cell):
    src = ''.join(cell.get('source', []))
    if not src.strip():
        return
    for line in src.split('\n'):
        stripped = line.strip()
        if stripped.startswith('#'):
            level = len(stripped) - len(stripped.lstrip('#'))
            heading_text = stripped.lstrip('#').strip()
            level = min(max(level, 1), 4)
            doc.add_heading(heading_text, level=level + 1)
        elif stripped:
            add_normal_text(doc, stripped)


def process_code_cell(doc, cell, nb_label, cell_idx):
    src = ''.join(cell.get('source', []))
    add_code_block(doc, src)

    for out in cell.get('outputs', []):
        otype = out.get('output_type')
        try:
            if otype == 'stream':
                text = ''.join(out.get('text', []))
                add_output_text(doc, text)
            elif otype in ('display_data', 'execute_result'):
                data = out.get('data', {})
                if 'image/png' in data:
                    img_data = data['image/png']
                    if isinstance(img_data, list):
                        img_data = ''.join(img_data)
                    img_bytes = base64.b64decode(img_data)
                    caption = f'{nb_label} - cell {cell_idx} output image'
                    add_image_from_bytes(doc, img_bytes, caption=caption)
                elif 'text/plain' in data:
                    text = data['text/plain']
                    if isinstance(text, list):
                        text = ''.join(text)
                    add_output_text(doc, text)
            elif otype == 'error':
                ename = out.get('ename', '')
                evalue = out.get('evalue', '')
                add_output_text(doc, f'[ERROR OUTPUT] {ename}: {evalue}')
        except Exception as e:
            add_normal_text(doc, f'[Could not process output #{cell_idx}: {e}]')


def process_notebook(doc, nb_path, nb_label):
    print(f'  Processing {nb_path} ...')
    doc.add_heading(nb_label, level=1)

    if not os.path.exists(nb_path):
        add_normal_text(doc, f'[NOTEBOOK NOT FOUND: {nb_path}]')
        print(f'    NOT FOUND, skipped.')
        return

    try:
        with open(nb_path, 'r', encoding='utf-8') as f:
            nb = json.load(f)
    except Exception as e:
        add_normal_text(doc, f'[Could not load notebook {nb_path}: {e}]')
        print(f'    ERROR loading: {e}')
        return

    cells = nb.get('cells', [])
    for idx, cell in enumerate(cells):
        ctype = cell.get('cell_type')
        try:
            if ctype == 'markdown':
                process_markdown_cell(doc, cell)
            elif ctype == 'code':
                process_code_cell(doc, cell, nb_label, idx)
        except Exception as e:
            add_normal_text(doc, f'[Could not process cell #{idx}: {e}]')
            print(f'    WARNING: cell #{idx} failed: {e}')

    print(f'    Done ({len(cells)} cells).')


# ----------------------------------------------------------------------------
# Section A: Benchmark metrics tables
# ----------------------------------------------------------------------------

def add_section_a(doc):
    doc.add_heading('Section A: Complete Metrics Tables', level=1)
    for pde in PDES:
        path = f'results/{pde}/benchmark_metrics.npy'
        doc.add_heading(f'{PDE_LABELS[pde]} - Benchmark Metrics', level=2)
        if not os.path.exists(path):
            add_normal_text(doc, f'[MISSING: {path}]')
            print(f'  Section A: MISSING {path}')
            continue
        try:
            metrics = np.load(path, allow_pickle=True).item()
            rows = []
            for model, m in metrics.items():
                rows.append([
                    model,
                    f"{m.get('l2', float('nan')):.6f}",
                    f"{m.get('max_error', float('nan')):.6f}",
                    f"{m.get('mae', float('nan')):.6f}",
                    f"{m.get('rmse', float('nan')):.6f}",
                ])
            make_table(doc, ['Model', 'Rel L2', 'Max Error', 'MAE', 'RMSE'], rows)
            print(f'  Section A: {pde} OK ({len(rows)} rows)')
        except Exception as e:
            add_normal_text(doc, f'[Could not load {path}: {e}]')
            print(f'  Section A: ERROR {pde}: {e}')


# ----------------------------------------------------------------------------
# Section B: Noise study table
# ----------------------------------------------------------------------------

def add_section_b(doc):
    doc.add_heading('Section B: Noise Study Table', level=1)
    rows = []
    for pde in PDES:
        path = f'results/{pde}/noise_study_metrics.npy'
        if not os.path.exists(path):
            add_normal_text(doc, f'[MISSING: {path}]')
            print(f'  Section B: MISSING {path}')
            continue
        try:
            data = np.load(path, allow_pickle=True).item()
            for noise_level, models in data.items():
                for model, m in models.items():
                    rows.append([
                        PDE_LABELS[pde],
                        noise_level,
                        model,
                        f"{m.get('l2', float('nan')):.6f}",
                        f"{m.get('max_error', float('nan')):.6f}",
                        f"{m.get('mae', float('nan')):.6f}",
                        f"{m.get('rmse', float('nan')):.6f}",
                    ])
            print(f'  Section B: {pde} OK')
        except Exception as e:
            add_normal_text(doc, f'[Could not load {path}: {e}]')
            print(f'  Section B: ERROR {pde}: {e}')

    if rows:
        make_table(doc, ['PDE', 'Noise Level', 'Model', 'Rel L2', 'Max Error', 'MAE', 'RMSE'], rows)
    else:
        add_normal_text(doc, '[No noise study data available]')


# ----------------------------------------------------------------------------
# Section C: Ablation table
# ----------------------------------------------------------------------------

def add_section_c(doc):
    doc.add_heading('Section C: Ablation Table', level=1)
    path = 'results/burgers/ablation_metrics.npy'
    if not os.path.exists(path):
        add_normal_text(doc, f'[MISSING: {path}]')
        print(f'  Section C: MISSING {path}')
        return
    try:
        metrics = np.load(path, allow_pickle=True).item()
        rows = []
        for model, m in metrics.items():
            rows.append([
                model,
                f"{m.get('l2', float('nan')):.6f}",
                f"{m.get('max_error', float('nan')):.6f}",
                f"{m.get('mae', float('nan')):.6f}",
                f"{m.get('rmse', float('nan')):.6f}",
            ])
        make_table(doc, ['Model', 'Rel L2', 'Max Error', 'MAE', 'RMSE'], rows)
        print(f'  Section C: OK ({len(rows)} rows)')
    except Exception as e:
        add_normal_text(doc, f'[Could not load {path}: {e}]')
        print(f'  Section C: ERROR: {e}')


# ----------------------------------------------------------------------------
# Section D: Epoch training data from CSVs
# ----------------------------------------------------------------------------

def add_section_d(doc):
    doc.add_heading('Section D: Epoch Training Data', level=1)
    for pde in PDES:
        fig_dir = f'figures/{pde}'
        if not os.path.isdir(fig_dir):
            continue
        csv_files = sorted(f for f in os.listdir(fig_dir) if f.endswith('_training.csv'))
        for csv_name in csv_files:
            csv_path = os.path.join(fig_dir, csv_name)
            exp_name = csv_name.replace('_training.csv', '')
            doc.add_heading(f'{PDE_LABELS[pde]} - {exp_name}', level=2)
            try:
                with open(csv_path, 'r', newline='') as f:
                    reader = csv.reader(f)
                    header = next(reader)
                    all_rows = list(reader)

                # columns: epoch, total, ic, bc, pde[, lambda_ic, lambda_bc, lambda_pde, stage]
                idx_map = {name: i for i, name in enumerate(header)}
                sampled = [r for r in all_rows if int(r[idx_map['epoch']]) % 500 == 0]

                table_rows = []
                for r in sampled:
                    table_rows.append([
                        r[idx_map['epoch']],
                        f"{float(r[idx_map['total']]):.6f}",
                        f"{float(r[idx_map['ic']]):.6f}",
                        f"{float(r[idx_map['bc']]):.6f}",
                        f"{float(r[idx_map['pde']]):.6f}",
                    ])
                make_table(doc, ['Epoch', 'Total Loss', 'IC Loss', 'BC Loss', 'PDE Loss'], table_rows)
                print(f'  Section D: {csv_path} OK ({len(table_rows)} sampled rows / {len(all_rows)} total)')
            except Exception as e:
                add_normal_text(doc, f'[Could not load {csv_path}: {e}]')
                print(f'  Section D: ERROR {csv_path}: {e}')


# ----------------------------------------------------------------------------
# Main
# ----------------------------------------------------------------------------

def main():
    print('Generating AC_PINN_Results.docx ...')
    doc = Document()

    # Default style tweaks
    style = doc.styles['Normal']
    style.font.name = TEXT_FONT
    style.font.size = TEXT_SIZE

    title = doc.add_heading('AC-PINN: Complete Results Document', level=0)
    add_normal_text(doc, 'Authors: Suyash Vasal Jain, Nishita Raghvendra')
    add_normal_text(doc, 'This document is a complete printout of all 7 executed notebooks, '
                         'plus metrics tables, ablation results, noise study results, and '
                         'sampled epoch training data.')
    doc.add_page_break()

    print('\nProcessing notebooks...')
    for nb_path, nb_label in NOTEBOOKS:
        try:
            process_notebook(doc, nb_path, nb_label)
        except Exception as e:
            add_normal_text(doc, f'[FATAL error processing {nb_path}: {e}]')
            print(f'  FATAL error on {nb_path}: {e}')
        doc.add_page_break()

    print('\nAdding additional sections...')
    try:
        add_section_a(doc)
    except Exception as e:
        add_normal_text(doc, f'[Section A failed: {e}]')
        print(f'  Section A FATAL: {e}')
    doc.add_page_break()

    try:
        add_section_b(doc)
    except Exception as e:
        add_normal_text(doc, f'[Section B failed: {e}]')
        print(f'  Section B FATAL: {e}')
    doc.add_page_break()

    try:
        add_section_c(doc)
    except Exception as e:
        add_normal_text(doc, f'[Section C failed: {e}]')
        print(f'  Section C FATAL: {e}')
    doc.add_page_break()

    try:
        add_section_d(doc)
    except Exception as e:
        add_normal_text(doc, f'[Section D failed: {e}]')
        print(f'  Section D FATAL: {e}')

    print(f'\nSaving document to {OUT_PATH} ...')
    doc.save(OUT_PATH)

    size_bytes = os.path.getsize(OUT_PATH)
    size_mb = size_bytes / (1024 * 1024)
    print(f'\nDone. {OUT_PATH} size: {size_bytes:,} bytes ({size_mb:.2f} MB)')


if __name__ == '__main__':
    main()
